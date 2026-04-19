import io
import logging
import re

from docx import Document

from backend.core.parsers.base import BaseParser, DocumentSection, ParsedDocument

logger = logging.getLogger(__name__)

# Word heading styles
_HEADING_STYLES = {"heading 1", "heading 2", "heading 3", "heading 4", "title"}

# Heuristic: all-caps or numbered lines as fallback section detection
_SECTION_RE = re.compile(
    r"^(\d+\.\s+[A-Z]|ARTICLE\s+[IVXLCDM\d]+|SECTION\s+\d+|[A-Z][A-Z\s]{4,}$)"
)


def _is_heading(para) -> bool:
    """Return True if the paragraph is a heading by style or heuristic."""
    style_name = para.style.name.lower() if para.style and para.style.name else ""
    if style_name in _HEADING_STYLES:
        return True
    text = para.text.strip()
    return bool(text and _SECTION_RE.match(text))


class DOCXParser(BaseParser):
    """Parse DOCX files using python-docx."""

    async def parse(self, file_bytes: bytes, filename: str) -> ParsedDocument:
        """Parse a DOCX into a ParsedDocument.

        Args:
            file_bytes: Raw DOCX bytes.
            filename: Original filename.

        Returns:
            ParsedDocument with extracted text and sections.

        Raises:
            ValueError: If the file cannot be parsed.
        """
        try:
            doc = Document(io.BytesIO(file_bytes))
        except Exception as exc:
            raise ValueError(f"Cannot open DOCX '{filename}': {exc}") from exc

        # Extract metadata from core properties
        props = doc.core_properties
        metadata: dict = {}
        if props.author:
            metadata["author"] = props.author
        if props.created:
            metadata["created"] = props.created.isoformat()
        if props.modified:
            metadata["modified"] = props.modified.isoformat()
        if props.title:
            metadata["title"] = props.title

        sections: list[DocumentSection] = []
        full_lines: list[str] = []
        char_offset = 0

        current_title: str | None = None
        current_lines: list[str] = []
        current_start: int = 0

        def _flush_section() -> None:
            nonlocal current_title, current_lines, current_start
            text = "\n".join(current_lines).strip()
            if text:
                sections.append(
                    DocumentSection(
                        text=text,
                        section_title=current_title,
                        char_offset_start=current_start,
                        char_offset_end=current_start + len(text),
                    )
                )
            current_lines = []

        for para in doc.paragraphs:
            text = para.text
            if not text.strip():
                continue

            if _is_heading(para):
                _flush_section()
                current_title = text.strip()
                current_start = char_offset
            else:
                current_lines.append(text)
                char_offset += len(text) + 1

            full_lines.append(text)

        _flush_section()

        full_text = "\n".join(full_lines)

        # DOCX has no native page count — estimate from section breaks
        section_count = len(doc.sections)

        logger.debug(
            "Parsed DOCX '%s': ~%d sections, %d clauses", filename, section_count, len(sections)
        )

        return ParsedDocument(
            filename=filename,
            file_type="docx",
            full_text=full_text,
            sections=sections,
            page_count=section_count or 1,
            metadata=metadata,
        )
